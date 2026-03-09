#!/usr/bin/env python3
"""
generate_docx.py — 마크다운 파일을 MS Word(.docx)로 변환하며 이미지를 삽입합니다.

사용법:
  python gateway/tools/generate_docx.py \
    --input output/{작품명}/proposal-client-{작품명}.md \
    --output output/{작품명}/proposal-client-{작품명}.docx \
    --image-base output/{작품명}/images
"""

import argparse
import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_heading_style(paragraph, level):
    """헤딩 색상 및 폰트 설정"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    if level == 1:
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x3C, 0x78)


def add_horizontal_rule(doc):
    """구분선 추가"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 40)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(9)


def resolve_image_path(img_ref, image_base, input_dir):
    """이미지 경로를 절대 경로로 변환"""
    # 경로 후보들
    candidates = [
        img_ref,                                          # 그대로
        os.path.join(image_base, os.path.basename(img_ref)),  # image_base 하위
        os.path.join(input_dir, img_ref),                # 입력 파일 상대 경로
        os.path.join(os.getcwd(), img_ref),              # 작업 디렉토리 기준
    ]
    for c in candidates:
        if c and os.path.isfile(c):
            return c
    return None


def parse_table(lines):
    """마크다운 테이블 파싱 → (headers, rows)"""
    headers = []
    rows = []
    for line in lines:
        if re.match(r'^\s*\|[-| :]+\|\s*$', line):
            continue  # 구분선 행 건너뜀
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def add_table_to_doc(doc, headers, rows):
    """Word 테이블 추가"""
    col_count = max(len(headers), max((len(r) for r in rows), default=0))
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'

    # 헤더 행
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers[:col_count]):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
        # 헤더 배경색
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E8F0FE')
        tcPr.append(shd)

    # 데이터 행
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row[:col_count]):
            row_cells[ci].text = cell_text

    doc.add_paragraph()  # 테이블 후 공백


def apply_inline_formatting(paragraph, text):
    """볼드(**), 이탤릭(*), 인라인코드(`) 처리"""
    # 패턴: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        # 앞 일반 텍스트
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        full = m.group(0)
        if full.startswith('**'):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif full.startswith('*'):
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif full.startswith('`'):
            run = paragraph.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def convert_md_to_docx(input_path, output_path, image_base):
    doc = Document()

    # 기본 폰트 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)

    input_dir = str(Path(input_path).parent)

    with open(input_path, encoding='utf-8') as f:
        content = f.read()

    lines = content.splitlines()
    i = 0
    table_buffer = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # 테이블 처리
        is_table_line = bool(re.match(r'^\s*\|', line))
        if is_table_line:
            table_buffer.append(line)
            in_table = True
            i += 1
            continue
        elif in_table:
            headers, rows = parse_table(table_buffer)
            if headers:
                add_table_to_doc(doc, headers, rows)
            table_buffer = []
            in_table = False
            # 현재 줄은 다시 처리

        # 빈 줄
        if not line.strip():
            i += 1
            continue

        # 구분선
        if re.match(r'^---+$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # 이미지 ![alt](path) 또는 [이미지 삽입: `path`]
        img_match = re.search(r'!\[.*?\]\((.+?)\)', line) or re.search(r'\[이미지 삽입[^\]]*[:`]\s*([^\]`]+)[`\]]', line)
        if img_match:
            img_ref = img_match.group(1).strip().strip('`').strip()
            img_path = resolve_image_path(img_ref, image_base, input_dir)
            if img_path:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                except Exception as e:
                    doc.add_paragraph(f'[이미지 삽입 실패: {img_ref} — {e}]')
            else:
                doc.add_paragraph(f'[이미지: {img_ref}]')
            i += 1
            continue

        # 헤딩
        h_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            # blockquote 제거
            text = re.sub(r'^>\s*', '', text)
            word_level = min(level, 4)
            p = doc.add_heading(text, level=word_level)
            set_heading_style(p, word_level)
            i += 1
            continue

        # 인용문 blockquote
        if line.startswith('>'):
            text = re.sub(r'^>\s*', '', line)
            p = doc.add_paragraph(style='Quote') if 'Quote' in [s.name for s in doc.styles] else doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            apply_inline_formatting(p, text)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # 불릿 리스트
        if re.match(r'^(\s*)[-*+]\s+', line):
            indent = len(re.match(r'^(\s*)', line).group(1))
            text = re.sub(r'^\s*[-*+]\s+', '', line)
            level = indent // 2
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
            apply_inline_formatting(p, text)
            i += 1
            continue

        # 번호 리스트
        if re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, text)
            i += 1
            continue

        # 일반 텍스트 (주석 제외)
        if line.startswith('<!--') or line.startswith('>>'):
            i += 1
            continue

        p = doc.add_paragraph()
        apply_inline_formatting(p, line)
        i += 1

    # 남은 테이블 처리
    if table_buffer:
        headers, rows = parse_table(table_buffer)
        if headers:
            add_table_to_doc(doc, headers, rows)

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    print(f"✅ Word 문서 저장 완료: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='마크다운 → Word(.docx) 변환 (이미지 삽입)')
    parser.add_argument('--input', required=True, help='입력 마크다운 파일 경로')
    parser.add_argument('--output', required=True, help='출력 .docx 파일 경로')
    parser.add_argument('--image-base', default='', help='이미지 기본 디렉토리 경로')
    args = parser.parse_args()

    if not os.path.isfile(args.input):
        print(f"❌ 입력 파일 없음: {args.input}")
        exit(1)

    convert_md_to_docx(args.input, args.output, args.image_base)


if __name__ == '__main__':
    main()
